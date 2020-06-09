import os
import requests
import docx
from lxml import etree
from fake_useragent import UserAgent


def write_doc(file_path, content_list):
    if os.path.isfile(file_path):
        doc = docx.Document(file_path)
    else:
        doc = docx.Document()

    for content in content_list:
        doc.add_paragraph(content)
    doc.save(file_path)



ua = UserAgent()
# base_url = "https://www.idnes.cz/sport/archiv/{}"
base_url = "https://www.blic.rs/"


def get_data():
    headers = ua.chrome
    response = requests.get(base_url, headers={"User_Agent": ua.chrome})
    # print(response.text)
    tree = etree.HTML(response.text)
    article_list = tree.xpath('//*[@id="top"]/div[4]/div/nav/ul/li/a/@href')
    print(article_list)
    for article in article_list[1:2]:  # 类型
        article = article.lower()
        print(article)
        type_url = f"https://www.blic.rs/{article}"
        print(type_url)
        article_type_list = requests.get(
            type_url, headers={"User_Agent": ua.chrome})
        res = article_type_list.text
        tree1 = etree.HTML(res)
        page_data = tree1.xpath(
            '//div[@class="pagination__list"]/ul/li/a/@href')[-2]
        page_res = page_data[:8]
        page = page_data[8:]
        print(f"{article}共有{page}页")
        for i in range(1, int(page)+1):  # 遍历每一页
            print(f"{article}类型，第{i}页", "*"*20)
            url = type_url + page_res+str(i)
            article_list = requests.get(
                url, headers={"User_Agent": ua.chrome}).text
            tree1 = etree.HTML(article_list)
            page_detail_url_list = tree1.xpath(
                '/html/body/main/div/section/div[2]/section/article/div/h3/a/@href')
            print(page_detail_url_list)
            for detail_url in page_detail_url_list:
                for i in range(5):
                    try:
                        article_detail = requests.get(
                            detail_url, headers={"User_Agent": ua.chrome})
                        if article_detail.status_code == 200:
                            break
                    except requests.exceptions.ProxyError:
                        continue
                else:
                    print(f"{detail_url}\t\t失败!!!")
                    continue
                tree2 = etree.HTML(article_detail.text)
                content = tree2.xpath(
                    '/html/body/main/div/article/div/div/p/text()')

                content = "".join(content).split(".")
                print(content)
                for index, item in enumerate(content):
                    if not item:
                        continue
                    content[index] = item + "."

                try:
                    write_doc("D:\stardata\小语种爬虫\赛尔/" +
                              "res_saier2.docx", content)
                except FileNotFoundError:
                    print(f"{detail_url}\t\t出错了!!!")
                    continue
                print(f"{detail_url}\t\t完成!!!")


if __name__ == '__main__':
    get_data()
