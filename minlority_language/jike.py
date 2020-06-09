import os
import requests
import docx
from lxml import etree
from fake_useragent import UserAgent


def write_doc(file_path, content_list):
    if os.path.isfile(file_path):
        doc = docx.Document(file_path)
    else:
        doc = docx.Document()

    for content in content_list:
        doc.add_paragraph(content)
    doc.save(file_path)


ua = UserAgent()
# base_url = "https://www.idnes.cz/sport/archiv/{}"
base_url = "https://www.idnes.cz/finance/archiv/{}"
start = int(input("起始页码>>>"))
end = int(input("终止页码>>>"))
for num in range(start, end + 1):
    print(f"第{num}页")
    headers = ua.chrome
    response = requests.get(base_url.format(str(num)),
                            headers={"User_Agent": ua.chrome})
    tree = etree.HTML(response.text)
    article_list = tree.xpath('//div[@class="art"]/a/@href')
    for article in article_list:
        article_url = article
        print(article_url)
        for i in range(5):
            try:
                article_detail = requests.get(
                    article_url, headers={"User_Agent": ua.chrome})
                if article_detail.status_code == 200:
                    break
            except requests.exceptions.ProxyError:
                continue
        else:
            print(f"{article_url}\t\t失败!!!")
            continue
        article_tree = etree.HTML(article_detail.text)
        content = article_tree.xpath('//div[@class="bbtext"]/p//text()')
        content = "".join(content).split(".")
        print(content)
        for index, item in enumerate(content):
            if not item:
                continue
            content[index] = item + "."

        try:
            write_doc("D:\stardata\小语种爬虫\捷克语/" + "res3.docx", content)
        except FileNotFoundError:
            print(f"{article_url}\t\t出错了!!!")
            continue
        print(f"{article_url}\t\t完成!!!")
    print(f"第{num}页完成!!!")
