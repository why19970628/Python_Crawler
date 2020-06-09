import os
import requests
import docx
from lxml import etree
from fake_useragent import UserAgent
import json


def write_doc(file_path, content_list):
    if os.path.isfile(file_path):
        doc = docx.Document(file_path)
    else:
        doc = docx.Document()

    for content in content_list:
        doc.add_paragraph(content)
    doc.save(file_path)


# 6
ua = UserAgent()
# base_url = "https://www.idnes.cz/sport/archiv/{}"
base_url = "https://www.seznam.cz/api/v2/graphql"

headers = ua.chrome
payloadData = {
    "operationName": "fetchTimelineWithUser",
    "query": "\n                query fetchTimelineWithUser($cursor: PostsCursor, $ids: [ID!], $type: [PostTypeEnum]) {\n                    posts(cursor: $cursor, ids: $ids, type: $type) {\n                        abVariant\n                        cursor\n                        requestId\n                        items {\n                            ...timelinePost\n                        }\n                    }\n                    \n                }\n\n                \n    fragment timelinePost on Post {\n        id\n        type {\n            name\n        }\n        title\n        postImage {\n            url\n        }\n        description\n        link\n        linkObscured # link pres trampolinu\n        source {\n            id\n            site {\n                title\n                hostname\n                iconURL\n                isSeznam\n                onHTTPS\n            }\n        }\n        alg\n        createTime\n        meta {\n            like {\n                total\n                currentUser\n            }\n            video {\n                spl\n            }\n            gallery {\n                allImagesUrl\n                images {\n                    title\n                    squareImage\n                    wideImage\n                }\n            }\n            videoportal {\n                service {\n                    favs\n                    isFav\n                    iconURL\n                    name\n                    url\n                    tagId\n                }\n            }\n        }\n        geo {\n            latitude\n        }\n        tags {\n            id\n            label\n            slug\n            typeName\n            relation\n            meta {\n                show\n                vpId\n            }\n            followed {\n                since\n            }\n            regUserCount\n        }\n    }\n\n            ",
    "variables": {
        "cursor": "H4sIAAAAAAAA_1Sauc5nuXHFcz9Gxxqg9mXiMQxlyhR_2x8wrMTqUSTo3Y3bFqbrFx6wLlks1nJYvP_89v63t4__-dt_f__9z799__brP7-Ziln-ku_vn695_6jPL-0W0Q-T_gj79uvvf__H15--WYdM_JIfLV-tL8tpfxfR-pLPr9fbEeuWR8x7Pr5e-ZLXl4j2e9jnG8RG_1_sfd_k7SNe9oh91ZeOXLH9t9jL4qWvtnzEtj4_dH-KZYw_YvHxNS__eNfXD93eVf31ccX2h27xqbFvqx_68YiVvvXn10-x0vwhVvpKG_G3z0e3in77ehnEfixaJqNTr5j8YZBPe1kdsTF7xLr1a98-I01_zDZRH2cLHfNDbN6_Xi7xqo94xNbL3v9YNLtHdn_CDfM6cHLtwqkL1ycujLzfbq5eWFhoVxMwMPPuz6lGRDQAuy80X8DCt-5zYRSmKs8LuzHVlAAO1t3CzAutnjgAzCusmhBW6KxmVw0NuWroDITX9UBTfGsN0z3-c-HsHXVRBYTp3BTQrQDTARtTRV5r-GLmUMwcofdQImHYKFgy2hRwrxoxe9fNkw0e6H6FMwxw_K5booAed90KdUAoWakGCCWrBKMd17A1ME4LXLQNztAeQghhr7uFjrmn3yUKqPi24YS9jm93oOTuXWgMLjre93wn4AxTctedRoROx113Jq4lZ-t-u8wba363sN4YTay7dR1YRSQAb3SriENYFaN6133gAt5U8KQNCNtJqg-sBGxoZYt1XTGVD0ZDoEZcN1OJxMwx2GAKhGsAJ7HuXidUlRuwqjp-YSTh3pk1-86sdROyKnKd6uCMdArCSzUWh6KLQzGF6czqKmmB7Vssvm14jrXi254EvAlZ3eCE7jeLqnthNLF9z5vN1BOH4tXXr3zhKr439T315xontK4aYTiUCLk7isb5PjzlwjUstH5HU26wa-pNfZomEHasmw5rpMMaieyt6YmZkcw1szHzwBo51HkQODkInBL4RuklGFraAwj3LsMJlkHnqktstKodEA5cDSVrEM61N6lqqxsg1Gjda5w2mK4NMdiBbNaBbNYJ3-iC6Xrg7U8UAiKOxi8L0nE44SS8bhILTV2S8EBMNYqFxq4l124pf3KMAkKrZW5f5vYN-NVGYLTh_NtUoy-x0V0E-8O2f0ITEQe8XNREbz0yyXsKJkjIJnkzoT3M5sK-RM6kr-eY7CVFpnrZiGkU4Y0j00yM5g0cs4AaljdwzOrmZ7O-hcDcbiyYI9eZg0Ca-7nFPBDrOhKjeQYWyuEozshLMXPhjLwGwn0JlXk3RkFrzSex7jR2NIN1Z6-dQ24hsJAbKRaKMwqDrcJwZGE39VkgmVsge1s4HCn8cmCLcMycsFVUYeZabGETC-1NBQ-8G0yBM6TdmmKJ_Gxp0DnB2y0dOidovKXf24QlaLxlwDcyHN9mQDix_USZsNJbNK3UOHoLgZViv2U3yVg5sko5tCqPAUTsVylGG25WjUOp20kYa3B-a8EWWmD2VsR-o6ZYO0zXDtN1IDQaNcUaNcU69zpDNzxnBME-KCI20YAZEE547BRnrkuZbNAreG7_V3gV-10wc1tFpKwpvjW7R7aFQ9lCYlzWhS3E_u71Zxe5ZneRq5WLXbO72LWzi19XcUERccm7rksJFqrCaBu02ptUXeXGkatdj3WNm-tcywH70kvXuafguonRxUIGZu4m11XcTBrQEjAh7LDGUwkujBtWbihebpn4FtXKDdXKDVcAt8WRud7Ggrve-uset0q6I5m79-UM7kjm7sjeHqIFCGsEqLgHqLgHUr0HsrdH4EADbRYPZGCPhD9HLaZC18WjFxBE3VOgc4J7e_q90j7sA6Poungix_pTKQAXU6Fb6yX3gvDABRwIKxyptDFqt0x4uQkhvvXEt2jveIVCOAwzxyVFXumYCsncC01jL9wXvBp-VQ1_rkYKqhEI4zbhtbBky71beUtgVAWjaCl7G7JK4-rhjT6St8En2xFW7bBke-FbdJm8A_m5Aw7cqRhFafNGLfNGl8m7kVS7C1o1YqHHodXAJ3saaqDB7j33nuK9AiX3kjEfXKZ80OzykcuCfHDV8jFY8mEcF_olRf5kjgsDlpxArptEIRhc4nxQrH1QnX1QnX3aoFXDRadhuhl47KDJ5rOI7llYY-WyIF9B7C9Zwapi1GCcNZTFJUlYPGT44iHDFw8ZvmQUG8jAi06gL26pvhlQA7dUX_KNLXjdoovo28iTO3CznYVWeG3xBXUJkZv6QsBkQvTaKsRu6guxa40QNCRDfDFz3FwXgnZlSAa-zcVCuKSHlELJKiyEN6CQSaw7g9G9hSBkA8JbWBfsK1RuXQi1e_qhIEWhdWM_FKQoFBft0L3lOHQVU6FNGroFrRZaGS7pYXI76mF6M0MY2FcYWqxhdjPDAxdwMBWaqGHRmKoM35ZDq1YshBexsMYp2AiEwfrCBRt0CYwqNuh2y0Q47uzhbhh17MhDMHPACT2xX0_4pOPpLRwkMJ4SC3ibt-GDkPRBSPrguB3PduEgkBFyC0E8IQqI0w_BcQeaAxFgmxGB7Udiv5E47siEGonQiMJxZ93s_ZBrBxwK37IY2dh-goxFzi18kbMYXcO3m5gZ_wZE4t-AKLweRsm91EThuTBKb5WMQg85ypyjgYXcCRFHhQ5zFJ78okDVotBhjmqcfg32W8P9gkFFgUFFLa2BHkUUXg-jFkm1wL6i0NmOBr-KluEoIrT18o1oNEOiDT7Z4GbRjnTdjqLZcSlENNhXdDaEQcaiCzHYeOKMbgRdD7e_3D4a7DHgVzHgVw-8Wo1iv2OIwUETJsbgSAO6FQO6Fc_9GBC2mkSJGXRsYgpuNg3fmEWhX_KcVTj_GvLkkqssHk9jE96-hfjdwilsI5y3kRi3cUbbcNGd-34UOwj2JbFZ_OsSu8ixS56zS2ugyxS7twClyKUfKXL5RgpYUIoKhPWeb4ouRtGhSkFLKsUKM4NfpXhiNBVapWEqNKxSwM1SKqFG31hIwQNKygjWHcXMUxAGVUvZm1RT5cZCKt67U8HcUvG8kmoGYbyYpOKBO9UXo3jRTs3BVKX4tm4cpYJuPRBa4b07DU22NLn0I01wvgYylqb3wSjNjPCmvjS8tqSBuaUZtm8uUAN_SaXhtSUNvC4Nry1paO6lgfWl4XklrRFH1gGtwNzSwNzSFpFiS0vi_6t0dALT8WSfrveSnm6whuPhJh1ELt0L3-JXrvRA4DiYWzqYWzpeedILkeKV0KoRKY4nofSBcRztynT8OZC-cDPfy2QyBK4SetlIht6e2wMXsDAVfkzN8FtwM_BjagbalRlxqVpGwnRRSBRRsMZDVi4c5MmYglbgsZmgppl4mcrEj22Z4KKZaCpm4m-0TL8FKBPUNBN9wsxYzFyoC4lOYCa4aGYv1Bgkxjz0sjT_oMD_-tO3v3_97z--vv_-59--_frtL9__8_tffvuvv3781X8xMRGz1ofl_es__i8AAP___b928_UuAAA",
        "limit": 100,
        "ids": 'null',
        "type": ["article", "video", "photo_gallery"]
    }
}

response = requests.post(base_url, data=json.dumps(
    payloadData), headers={"User_Agent": ua.chrome})
print(response.text)
data = response.text.replace("null", "None").replace(
    "false", "None").replace("true", "None")
data = eval(data).get('data').get("posts").get("items")
for d in data:
    link = d.get("link")
    print(link)
print(data)
article_list = tree.xpath('//div[@class="art"]/a/@href')
for article in article_list:
    article_url = article
    print(article_url)
    for i in range(5):
        try:
            article_detail = requests.get(
                article_url, headers={"User_Agent": ua.chrome})
            if article_detail.status_code == 200:
                break
        except requests.exceptions.ProxyError:
            continue
    else:
        print(f"{article_url}\t\t失败!!!")
        continue
    article_tree = etree.HTML(article_detail.text)
    content = article_tree.xpath('//div[@class="bbtext"]/p//text()')
    content = "".join(content).split(".")
    print(content)
    for index, item in enumerate(content):
        if not item:
            continue
        content[index] = item + "."

    try:
        write_doc("D:\stardata\小语种爬虫\捷克语/" + "res3.docx", content)
    except FileNotFoundError:
        print(f"{article_url}\t\t出错了!!!")
        continue
    print(f"{article_url}\t\t完成!!!")
# print(f"第{num}页完成!!!")
